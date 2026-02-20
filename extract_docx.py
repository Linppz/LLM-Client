import sys
import os

# Force UTF-8 output
os.environ['PYTHONIOENCODING'] = 'utf-8'

from docx import Document

doc = Document(r'C:\Users\lllpp\OneDrive\桌面\Week3_Prompt工程化学习计划(1).docx')

with open(r'C:\Users\lllpp\llm-client\day5_output.txt', 'w', encoding='utf-8') as f:
    f.write("=== PARAGRAPHS ===\n")
    for i, p in enumerate(doc.paragraphs):
        f.write(f"[{i}] {p.text}\n")
    
    f.write("\n=== TABLES ===\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\n--- Table {t_idx} ---\n")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text for cell in row.cells]
            f.write(f"Row {r_idx}: {' | '.join(cells)}\n")

print("Extraction complete")
